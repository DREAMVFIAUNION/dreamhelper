"""文档解析引擎 — 支持 PDF/DOCX/XLSX/CSV/MD/TXT/JSON

从上传的文件中提取纯文本内容，用于注入对话上下文。
"""

import io
import csv
import json
import logging
from typing import Optional

logger = logging.getLogger(__name__)

MAX_FILE_BYTES = 20 * 1024 * 1024  # 20MB

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt", ".md", ".csv", ".json"}


async def parse_document(file_bytes: bytes, filename: str) -> dict:
    """解析文档 → 提取纯文本

    Returns:
        {"success": bool, "text": str, "doc_type": str, "error": str | None}
    """
    if not file_bytes:
        return {"success": False, "text": "", "doc_type": "", "error": "空文件"}

    if len(file_bytes) > MAX_FILE_BYTES:
        return {
            "success": False, "text": "", "doc_type": "",
            "error": f"文件过大: {len(file_bytes) / 1024 / 1024:.1f}MB，上限 {MAX_FILE_BYTES // 1024 // 1024}MB",
        }

    ext = _get_ext(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        return {
            "success": False, "text": "", "doc_type": ext,
            "error": f"不支持的文件类型: {ext}。支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}",
        }

    try:
        if ext == ".pdf":
            text = _parse_pdf(file_bytes)
            doc_type = "pdf"
        elif ext == ".docx":
            text = _parse_docx(file_bytes)
            doc_type = "docx"
        elif ext == ".xlsx":
            text = _parse_xlsx(file_bytes)
            doc_type = "xlsx"
        elif ext == ".csv":
            text = _parse_csv(file_bytes)
            doc_type = "csv"
        elif ext == ".json":
            text = _parse_json(file_bytes)
            doc_type = "json"
        elif ext == ".md":
            text = file_bytes.decode("utf-8", errors="replace")
            doc_type = "markdown"
        else:  # .txt
            text = file_bytes.decode("utf-8", errors="replace")
            doc_type = "text"

        if not text.strip():
            return {"success": False, "text": "", "doc_type": doc_type, "error": "文档内容为空"}

        return {"success": True, "text": text.strip(), "doc_type": doc_type, "error": None}

    except Exception as e:
        logger.error("文档解析失败 [%s]: %s", filename, e, exc_info=True)
        return {"success": False, "text": "", "doc_type": ext, "error": f"解析失败: {str(e)[:200]}"}


def _get_ext(filename: str) -> str:
    parts = filename.rsplit(".", 1)
    return ("." + parts[-1].lower()) if len(parts) > 1 else ""


def _parse_pdf(data: bytes) -> str:
    """PDF → 文本 (PyMuPDF / fitz)"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return "[PDF 解析需要安装 PyMuPDF: pip install PyMuPDF]"

    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"--- 第 {i + 1} 页 ---\n{text.strip()}")
    doc.close()
    return "\n\n".join(pages) if pages else ""


def _parse_docx(data: bytes) -> str:
    """DOCX → 文本 (python-docx)"""
    try:
        from docx import Document
    except ImportError:
        return "[DOCX 解析需要安装 python-docx: pip install python-docx]"

    doc = Document(io.BytesIO(data))
    parts = []

    # 段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 表格 → Markdown 格式
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            # 在第一行后插入分隔线
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, header_sep)
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


def _parse_xlsx(data: bytes) -> str:
    """XLSX → Markdown 表格 (openpyxl)"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "[XLSX 解析需要安装 openpyxl: pip install openpyxl]"

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheets = []

    for ws in wb.worksheets:
        rows_data = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):  # 跳过全空行
                rows_data.append(cells)

        if not rows_data:
            continue

        # 转 Markdown 表格
        max_cols = max(len(r) for r in rows_data)
        lines = [f"### 工作表: {ws.title}"]

        for i, row in enumerate(rows_data):
            # 补齐列数
            padded = row + [""] * (max_cols - len(row))
            lines.append("| " + " | ".join(padded) + " |")
            if i == 0:
                lines.append("| " + " | ".join(["---"] * max_cols) + " |")

        sheets.append("\n".join(lines))

    wb.close()
    return "\n\n".join(sheets)


def _parse_csv(data: bytes) -> str:
    """CSV → Markdown 表格"""
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return ""

    max_cols = max(len(r) for r in rows)
    lines = []
    for i, row in enumerate(rows):
        padded = row + [""] * (max_cols - len(row))
        lines.append("| " + " | ".join(padded) + " |")
        if i == 0:
            lines.append("| " + " | ".join(["---"] * max_cols) + " |")

    return "\n".join(lines)


def _parse_json(data: bytes) -> str:
    """JSON → 格式化输出"""
    text = data.decode("utf-8", errors="replace")
    try:
        obj = json.loads(text)
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return text
